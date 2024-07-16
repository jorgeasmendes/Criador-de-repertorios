import pandas as pd
import numpy
import re
import tkinter as tk
from tkinter import ttk 
import baixando_cifras_letras
from unidecode import unidecode
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor

##Widgets personalizados
 #Entry com limite de caracteres
class EntryLimit(tk.Entry):
    def __init__(self, master, limit, width=10):
        self.limit_var=limit
        tk.Entry.__init__(self, master=master, width=width)
        self.var=tk.StringVar()
        self.var.trace_add(mode='write', callback=self.valida_entry)
        self['textvariable']=self.var
    def valida_entry(self, name, index, mode):
        if len(self.var.get())>0:
            self.var.set(self.var.get()[:self.limit_var])


  #Combobox com limite de caracteres e listbox acoplada para dar sugestão de preenchimento do campo, conforme o usuário digita.
class MeuComboBox():
    def __init__(self, master, limit, dados, width=10):
        self.dados=dados
        self.limit=limit
        self.width=width
        self.janela=master
        self.data=[]
        self.var_str=tk.StringVar()
        self.var_str.trace_add(mode='write', callback=self.valida)
    
    def pack(self): #função para dar visibilidade ao objeto
        self.frame=tk.Frame(self.janela)
        self.frame.pack()
        self.combo=ttk.Combobox(self.frame, width=self.width, values=self.dados,textvariable=self.var_str)
        self.combo.pack()
        self.list_box=tk.Listbox(self.frame, listvariable=self.data, height=0, width=self.width, selectmode='browse', activestyle='dotbox')
        self.list_box.bind('<<ListboxSelect>>',self.clica_lista)
                
    def clica_lista(self,evt): #função para o preenchimento do campo ao clicar na listbox
        try:
            selecionado=self.list_box.curselection()
            self.combo.set(self.list_box.get(selecionado))
            self.combo.focus()
        except:
            pass
        self.list_box.pack_forget()
    
    def valida(self,var,index,write):  #limitação do número de caracteres e surgimento das sugestões depreenchimento na listbox
        self.list_box.delete(0, tk.END)
        if len(self.var_str.get())==0:
            self.list_box.pack_forget()
            self.combo['values']=self.dados
        elif len(self.var_str.get())>0:
            self.combo.set(self.combo.get()[:self.limit]) #limite de carcteres
            data=[]
            indice_lista=0
            for i in self.dados:
                if i.lower().startswith(self.var_str.get().lower()):  #insere na listbox os valores que iniciam com os caracteres digitados
                    data.append(i)
                    self.list_box.pack()
                    self.list_box.insert(indice_lista, i)
                    indice_lista+=1
            if data == []:
                self.combo['values']=self.dados
                self.list_box.pack_forget()
            else:
                self.combo['values']=data

#Variáveis iniciais: valores padrão para preenchimento dos campos e nome do arquivo onde será armazenada a lista do repertório.
tonalidades = ['','C','Db','D','Eb','E','F','F#','Gb','G','Ab','A','Bb','B',
               'Cm','C#m','Dm','Ebm','Em','Fm','F#m','Gm','G#m','Am','A#m','Bbm','Bm']
andamentos = ['','1 - Muito lento','2 - Lento','3 - Médio','4 - Rápido','5 - Muito Rápido']
arquivo='repertorio.csv'


##início do programa e janela inicial de consulta
  #função para abrir csv e inserir dados na coluna ordenados do mais recente ao  mais antigo
def open():
    global df
    tabela.delete(*tabela.get_children())
    df=pd.read_csv(arquivo)
    df_exibicao=df.sort_index(ascending=False)
    tabela['column']=list(df_exibicao.columns)
    tabela['show']= 'headings'
    for col in tabela['column']:
        tabela.column(col,anchor='center', stretch=False, width=100)
        tabela.heading(col, text=col)
    linhas=df_exibicao.to_numpy().tolist()
    for linha in linhas:
        tabela.insert('','end',values=linha)

def cadastrar():
    janela_cadastro()

#deletar dado selecionado ao pressionar tecla delete
def delet(event):
    global df
    if tabela.selection() != ():
        selected_item = tabela.selection()[0]
        nome=tabela.item(selected_item,'values')[0]
        df=df.loc[df['nome'] != nome]
        df.to_csv('repertorio.csv', index=False)
        tabela.delete(selected_item)

#deletar dado selecionado ao clicar no botão delete
def deletar():
    global df
    if tabela.selection() != ():
        selected_item = tabela.selection()[0]
        nome=tabela.item(selected_item,'values')[0]
        df=df.loc[df['nome'] != nome]
        df.to_csv('repertorio.csv', index=False)
        tabela.delete(selected_item)
    

#janela inicial
def janela_consultar():
    global tabela
    janela_consulta=tk.Tk()
    janela_consulta.title('Consulta de músicas')
    janela_consulta.geometry('920x550')

    titulo=tk.Label(janela_consulta, text='CRIADOR DE REPERTÓRIOS')
    titulo.pack()

    menu=tk.Frame(janela_consulta)
    menu.pack()
    botao_cadastro=tk.Button(menu,text='Cadastrar Nova Música',command=cadastrar)
    botao_cadastro.grid(row=0, column= 0)
    botao_del=tk.Button(menu,text='Deletar',command=deletar)
    botao_del.grid(row=0, column= 1)

    tabela=ttk.Treeview(janela_consulta, selectmode='browse',height=20)
    tabela.pack(side ='left')
    tabela.heading('#0', text='\n')
    tabela.bind('<Delete>', delet)
    scrlbar_y = ttk.Scrollbar(janela_consulta, orient ="vertical", command = tabela.yview)
    scrlbar_y.pack(side ='right', fill ='y')
    tabela.configure(yscrollcommand = scrlbar_y.set)
    open() 
    janela_consulta.mainloop()



##--------------------------
##Janela de cadastro
#comando do botão salvar
def salvar_cadastro():
    global df
    if insert_nome.get().strip() == '':
        pop_up('vazio')
    else:
        lista_nomes=[]
        for nome in df['nome'].tolist():
            lista_nomes.append(nome.lower())

        if insert_nome.get().lower() in lista_nomes:
            pop_up('existe')

        else:
            novo_cadastro= {'nome':insert_nome.get(), 'tom':insert_tom.get(),'compositor':insert_compositor.entry.get(),
                            'artista':insert_artista.entry.get(),'genero':insert_genero.entry.get(),
                            'subgenero':insert_subgenero.entry.get(),'andamento':insert_andamento.get(),
                            'clima':insert_clima.entry.get(),'tema':insert_tema.entry.get()}
            
            df = df._append(novo_cadastro,ignore_index=True)
            df.to_csv(arquivo,index=False)
            tabela.insert('',index=0,values=[insert_nome.get(),insert_tom.get(),insert_compositor.entry.get(),insert_artista.entry.get(),
                                            insert_genero.entry.get(),insert_subgenero.entry.get(),insert_andamento.get(),
                                            insert_clima.entry.get(),insert_tema.entry.get()])
            pop_up('sucesso')

#popup de confirmação de salvamento ou erro caso a música já esteja cadastrada ou o campo de nome da música esteja vazio
def pop_up(situacao):
    def pop_up_exit(event): janela.grab_set()
    def enter(event): janelaok.destroy()   
    def botao_ok_press(): janelaok.destroy()   
    janelaok=tk.Toplevel()
    janelaok.grab_set()
    janelaok.bind('<Destroy>', pop_up_exit)
    janelaok.geometry('300x100') 
    janelaok.bind('<Return>', enter)
    janelaok.focus_force()
    if situacao == 'existe':
        janelaok.title('Erro no salvamento')
        mensagem=tk.Label(janelaok, text='Já existe uma música registrada com esse nome.\nTente novamente.')
        mensagem.pack()
    elif situacao =='vazio':
        janelaok.title('Erro no salvamento')
        mensagem=tk.Label(janelaok, text='Você precisa inserir o título da música')
        mensagem.pack()
    elif situacao == 'sucesso':
        janelaok.title('Música salva')
        mensagem=tk.Label(janelaok, text='Música salva com sucesso')
        mensagem.pack()

    botao_ok=tk.Button(janelaok, text='Ok', command=botao_ok_press)
    botao_ok.pack()
    janelaok.mainloop()

#janela de anexação de letra/cifra/ partitura
def anexar_arquivo():
    def pop_up_exit(event): janela.grab_set() 
    def abrir_arquivo():
        pass
    #baixar letra e cifra usando webscrap dos sites cifraclub.com.br e letras.com
    def baixar_letra(): 
        global tipo_documento
        texto.delete(1.0, tk.END)
        texto.insert(1.0, baixando_cifras_letras.baixar_letra(artista=unidecode(insert_artista.entry.get().lower().replace(' ','-')), musica=unidecode(insert_nome.get().lower().replace(' ','-'))))
        tipo_documento='letra'
    def baixar_cifra():
        global tipo_documento
        texto.delete(1.0, tk.END)
        texto.insert(1.0, baixando_cifras_letras.baixar_cifra(artista=unidecode(insert_artista.entry.get().lower().replace(' ','-')), musica=unidecode(insert_nome.get().lower().replace(' ','-')), semitons=int(modular_box.get())))
        tipo_documento='cifra'
    
    def anexar():    #configurações e salvamento do documento docx
        document = Document()

        titulo=document.add_paragraph().add_run(insert_nome.get())
        fonte_titulo=titulo.font
        fonte_titulo.name = 'Montserrat'
        fonte_titulo.size = Pt(25)
        fonte_titulo.bold=True
        #fonte_titulo.color.rgb= RGBColor(23, 54, 93)

        subtitulo=document.add_paragraph().add_run(f'{insert_compositor.entry.get()}\nTom: {insert_tom.get()}')
        fonte_subtitulo=subtitulo.font
        fonte_subtitulo.name = 'Calibri'
        fonte_subtitulo.size = Pt(14)
        fonte_subtitulo.bold=True

        corpo=document.add_paragraph().add_run(texto.get(1.0, tk.END))
        fonte=corpo.font
        fonte.name = 'Arial'
        document.add_page_break()
        if tipo_documento == 'cifra':
            fonte.size = Pt(12)
            document.save(f'Cifras/{insert_nome.get()}.docx')
        elif tipo_documento == 'letra':
            fonte.size = Pt(15)
            document.save(f'Letras/{insert_nome.get()}.docx')

    janela_arquivo=tk.Toplevel()
    janela_arquivo.grab_set()
    janela_arquivo.bind('<Destroy>', pop_up_exit)
    janela_arquivo.geometry('670x480') 
    janela_arquivo.focus_force()
    janela_arquivo.title('Anexar letra/cifra/partitura')
    titulo=tk.Label(janela_arquivo, text='Anexar letra/cifra/partitura')
    titulo.pack()
    menu=tk.Frame(janela_arquivo)
    menu.pack()
    botao_manualmente=tk.Button(menu,text='Selecionar arquivo',command=abrir_arquivo)
    botao_manualmente.grid(row=0, column= 0)
    botao_letra=tk.Button(menu,text='Baixar letra do Letras.com',command=baixar_letra)
    botao_letra.grid(row=0, column= 1)
    botao_cifra=tk.Button(menu,text='Baixar cifras do CifraClub',command=baixar_cifra)
    botao_cifra.grid(row=0, column= 2)

#botão para selecionar a modulação de tonalidade desejada para a cifra
    modular=tk.Frame(menu)
    modular.grid(row=0, column= 3)  
    modular_label=tk.Label(modular,text='Modular cifra\n(número de semitons ascendentes)')
    modular_label.pack()
    modular_box=ttk.Combobox(modular, values=[0,1,2,3,4,5,6,7,8,9,10,11,12], state='readonly',width=5)
    modular_box.set(0)
    modular_box.pack()

    botao_anexar=tk.Button(menu,text='ANEXAR',command=anexar)
    botao_anexar.grid(row=0, column= 4)

#janela para editar texto da cifra ou da letra
    texto=tk.Text(janela_arquivo)
    texto.pack(side='left')
    scrltext_y = ttk.Scrollbar(janela_arquivo, orient ="vertical", command = texto.yview)
    scrltext_y.pack(side ='right', fill ='y')
    texto.configure(yscrollcommand = scrltext_y.set)
    janela_arquivo.mainloop()

    
##janela para cadastrar novas músicas ou editar cadastro já salvo
def janela_cadastro():
    global janela, insert_nome, insert_tom, insert_compositor, insert_artista, insert_genero, insert_subgenero, insert_andamento, insert_clima, insert_tema
    janela=tk.Toplevel()
    janela.grab_set()
    janela.title('Cadastro de nova música')
    janela.geometry('500x500')
    janela.focus_force()



    label_nome=tk.Label(janela,text='Nome da música')
    label_nome.pack()
    insert_nome=EntryLimit(master=janela,width=55,limit=50)
    insert_nome.pack(pady=5)

    janela_tom_andamento=tk.Frame(janela)
    janela_tom_andamento.pack()
    label_tom=tk.Label(janela_tom_andamento,text='Tom')
    label_tom.grid(row=0, column=1, padx=30)
    insert_tom=ttk.Combobox(janela_tom_andamento, values=tonalidades, state='readonly',width=5)
    insert_tom.grid(row=1, column=1, padx=30)

    label_andamento=tk.Label(janela_tom_andamento,text='Andamento')
    label_andamento.grid(row=0, column=0, padx=10)
    insert_andamento=ttk.Combobox(janela_tom_andamento, values=andamentos, state='readonly')
    insert_andamento.grid(row=1, column=0, padx=10)

    janela_compositor=tk.Frame(janela)
    janela_compositor.pack()
    label_compositor=tk.Label(janela_compositor,text='Compositor')
    label_compositor.pack()
    insert_compositor=MeuComboBox(master=janela_compositor, limit=50, width=55, dados=df['compositor'].unique().tolist())
    insert_compositor.pack()
    

    janela_artista=tk.Frame(janela)
    janela_artista.pack()
    label_artista=tk.Label(janela_artista,text='Artista')
    label_artista.pack()
    insert_artista=MeuComboBox(master=janela_artista, limit=50, width=55, dados=df['artista'].unique().tolist())
    insert_artista.pack()

    janela_genero=tk.Frame(janela)
    janela_genero.pack()
    label_genero=tk.Label(janela_genero,text='Gênero')
    label_genero.pack()
    insert_genero=MeuComboBox(master=janela_genero, limit=50, width=55, dados=df['genero'].unique().tolist())
    insert_genero.pack()

    janela_subgenero=tk.Frame(janela)
    janela_subgenero.pack()
    label_subgenero=tk.Label(janela_subgenero,text='Sub-Gênero')
    label_subgenero.pack()
    insert_subgenero=MeuComboBox(master=janela_subgenero, limit=50, width=55, dados=df['subgenero'].unique().tolist())
    insert_subgenero.pack()


    janela_clima=tk.Frame(janela)
    janela_clima.pack()
    label_clima=tk.Label(janela_clima,text='Clima')
    label_clima.pack()
    insert_clima=MeuComboBox(master=janela_clima, limit=50, width=55, dados=df['clima'].unique().tolist())
    insert_clima.pack()

    janela_tema=tk.Frame(janela)
    janela_tema.pack()
    label_tema=tk.Label(janela_tema,text='Tema')
    label_tema.pack()
    insert_tema=MeuComboBox(master=janela_tema, limit=50, width=55, dados=df['tema'].unique().tolist())
    insert_tema.pack()

    arquivos=tk.Button(janela,text='Anexar letra/cifra/partitura',command=anexar_arquivo)
    arquivos.pack(pady=20)

    salvar=tk.Button(janela,text='Salvar nova música',command=salvar_cadastro)
    salvar.pack(pady=20)
    janela.mainloop()


#Ativar início do programa
janela_consultar()
        




